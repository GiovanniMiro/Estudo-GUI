import pandas as pd
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt


#----------------------------------------------------------------
root = Tk()
root.title("Gerador de Certificado")
#----------------------------------------------------------------
stilo = ttk.Style()
stilo.theme_use("alt")
stilo.configure(".", font=("Arial 10"), rowheight=30)

treeviewDados = ttk.Treeview(root, columns=(1, 2 , 3, 4, 5, 6), show="headings")

treeviewDados.column("1", anchor=CENTER)
treeviewDados.heading("1", text="CPF")

treeviewDados.column("2", anchor=CENTER)
treeviewDados.heading("2", text="Nome")

treeviewDados.column("3", anchor=CENTER)
treeviewDados.heading("3", text="RG")

treeviewDados.column("4", anchor=CENTER)
treeviewDados.heading("4", text="Data Início")

treeviewDados.column("5", anchor=CENTER)
treeviewDados.heading("5", text="Data Fim")

treeviewDados.column("6", anchor=CENTER)
treeviewDados.heading("6", text="Email")

treeviewDados.grid(row=4, column=0, columnspan=10, sticky="NSEW", pady=15)

def transferirTreeviewEntries(Event):

    item = treeviewDados.selection()

    for i in item:

        #Limpando os campos de entrada de dados.
        inserirCPF.delete(0, END)
        inserirNome.delete(0, END)
        inserirRG.delete(0, END)
        inserirDataInicio.delete(0, END)
        inserirDataFim.delete(0, END)
        inserirEmail.delete(0, END)

        inserirCPF.insert(0, treeviewDados.item(i, "values")[0])
        inserirNome.insert(0, treeviewDados.item(i, "values")[1])
        inserirRG.insert(0, treeviewDados.item(i, "values")[2])
        inserirDataInicio.insert(0, treeviewDados.item(i, "values")[3])
        inserirDataFim.insert(0, treeviewDados.item(i, "values")[4])
        inserirEmail.insert(0, treeviewDados.item(i, "values")[5])

treeviewDados.bind("<Double-1>", transferirTreeviewEntries)
#----------------------------------------------------------------------

dadosUsuarios = pd.read_excel("Dados.xlsx")

dadosUsuarios["Data Inicio"] = dadosUsuarios["Data Inicio"].astype(str)
dadosUsuarios["Data Fim"] = dadosUsuarios["Data Fim"].astype(str)

#-----------------------------------------------------------------------
for linha in range(len(dadosUsuarios)):

    dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
    dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
    dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

    dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

    dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
    dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
    dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

    dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

    treeviewDados.insert("","end",
                         values=(str(dadosUsuarios.iloc[linha, 0]),
                                 str(dadosUsuarios.iloc[linha, 1]),
                                 str(dadosUsuarios.iloc[linha, 2]),
                                 str(dataInicioTratada),
                                 str(dataFimTratada),
                                 str(dadosUsuarios.iloc[linha, 5])))

#---------------------------------------------------------------------

labelCPF = Label(text="CPF: ", font="Arial 12").grid(row=0, column=0, sticky="E", pady=5)
labelNome = Label(text="Nome: ", font="Arial 12").grid(row=0, column=3, sticky="E", pady=5)
labelRG = Label(text="RG: ", font="Arial 12").grid(row=0, column=6, sticky="E", pady=5)
labelDataInicio = Label(text="Data Inicio: ", font="Arial 12").grid(row=2, column=0, sticky="E", pady=5)
labelDataFim = Label(text="Data Fim: ", font="Arial 12").grid(row=2, column=3, sticky="E", pady=5)
labelEmail = Label(text="Email: ", font="Arial 12").grid(row=2, column=6, sticky="E", pady=5)

#----------------------------------------------------------------------

inserirCPF = Entry(font="Arial 12")
inserirCPF.grid(row=0, column=1, columnspan=2, sticky="W", pady=5)

inserirNome = Entry(font="Arial 12")
inserirNome.grid(row=0, column=4, columnspan=2, sticky="W", pady=5)

inserirRG = Entry(font="Arial 12")
inserirRG.grid(row=0, column=7, columnspan=2, sticky="W", pady=5)

inserirDataInicio = Entry(font="Arial 12")
inserirDataInicio.grid(row=2, column=1, columnspan=2, sticky="W", pady=5)

inserirDataFim = Entry(font="Arial 12")
inserirDataFim.grid(row=2, column=4, columnspan=2, sticky="W", pady=5)

inserirEmail = Entry(font="Arial 12")
inserirEmail.grid(row=2, column=7, columnspan=2, sticky="W", pady=5)

#---------------------------------------------------------------------

def filtrar_dados():

    print("Filtrando...")

    for linha in range(len(dadosUsuarios)):

        todasLinhas = treeviewDados.get_children()

        #Deletando todas as linhas da treeview
        treeviewDados.delete(*todasLinhas)

        #Se o CPF for vazio, não tiver informação, carrega todos os dados
        if inserirCPF.get() == "":

            #Mudando o texto do botão
            botaoPesquisar.config(text="Filtrar")

            for linha in range(len(dadosUsuarios)):

                dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
                dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                treeviewDados.insert("", "end",
                                     values=(str(dadosUsuarios.iloc[linha, 0]),
                                             str(dadosUsuarios.iloc[linha, 1]),
                                             str(dadosUsuarios.iloc[linha, 2]),
                                             str(dataInicioTratada),
                                             str(dataFimTratada),
                                             str(dadosUsuarios.iloc[linha, 5])))

        else:

            botaoPesquisar.config(text="Limpar Filtros")

            for linha in range(len(dadosUsuarios)):

                dataInicioAno = dadosUsuarios.iloc[linha, 3].split("-")[0]
                dataInicioMes = dadosUsuarios.iloc[linha, 3].split("-")[1]
                dataInicioDia = dadosUsuarios.iloc[linha, 3].split("-")[2]

                dataInicioTratada = dataInicioDia + "/" + dataInicioMes + "/" + dataInicioAno

                dataFimAno = dadosUsuarios.iloc[linha, 4].split("-")[0]
                dataFimMes = dadosUsuarios.iloc[linha, 4].split("-")[1]
                dataFimDia = dadosUsuarios.iloc[linha, 4].split("-")[2]

                dataFimTratada = dataFimDia + "/" + dataFimMes + "/" + dataFimAno

                #Verifico se o CPF do campo Entry é igual ao CPF da linha corrente do treeview
                if inserirCPF.get() == str(dadosUsuarios.iloc[linha, 0]):

                    treeviewDados.insert("", "end",
                                         values=(str(dadosUsuarios.iloc[linha, 0]),
                                                 str(dadosUsuarios.iloc[linha, 1]),
                                                 str(dadosUsuarios.iloc[linha, 2]),
                                                 str(dataInicioTratada),
                                                 str(dataFimTratada),
                                                 str(dadosUsuarios.iloc[linha, 5])))

def gerar_certificado():

    print("Gerando certificado...")

    #Abrindo o arquivo do Word
    arquivoWord = Document("Certificado.docx")

    #Configurando os estilos
    estilo = arquivoWord.styles["Normal"]

    #Pegamos os dados do aluno dos campo Entry
    nomeAluno = inserirNome.get()
    dataInicio = inserirDataInicio.get()
    dataFim = inserirDataFim.get()
    nomeInstrutor = "Nome do Instrutor"
    CPFAluno = inserirCPF.get()
    RGAluno = inserirRG.get()

    frase_parte1 = "concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
    frase_montada = f"{nomeAluno}, CPF: {CPFAluno}, RG: {RGAluno}, {frase_parte1} {dataInicio} a {dataFim}."

    for paragrafo in arquivoWord.paragraphs:

        if "@nome" in paragrafo.text:

            paragrafo.text = nomeAluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@DataFim" in paragrafo.text:

            paragrafo.text = frase_montada
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@instrutor" in paragrafo.text:

            paragrafo.text = f"{nomeInstrutor} - Instrutor"
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    #Montando o caminho + nome do certificado
    caminhoCertificadoGerado = nomeAluno + ".docx"

    #Salvando o certificado com o nome do aluno
    arquivoWord.save(caminhoCertificadoGerado)

    inserirCPF.delete(0, END)
    inserirNome.delete(0, END)
    inserirRG.delete(0, END)
    inserirDataInicio.delete(0, END)
    inserirDataFim.delete(0, END)
    inserirEmail.delete(0, END)

    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso!")

def gerar_certificado_massa():

    print("Gerando certificados em massa...")

    #Passo linha por linha
    for linha in treeviewDados.get_children():

        #Os valores da linha corrente em forma de colunas
        coluna = treeviewDados.item(linha)["values"]

        CPF_separado = coluna[0]
        nome_separado = coluna[1]
        RG_separado = coluna[2]
        dataInicio_separado = coluna[3]
        dataFim_separado = coluna[4]

        #Abrindo o arquivo do Word
        arquivoWord = Document("Certificado.docx")

        #Configurando os estilos
        estilo = arquivoWord.styles["Normal"]

        nomeInstrutor = "Nome do Instrutor"

        frase_parte1 = "concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de "
        frase_montada = f"{nome_separado}, CPF: {CPF_separado}, RG: {RG_separado}, {frase_parte1} {dataInicio_separado} a {dataFim_separado}."

        for paragrafo in arquivoWord.paragraphs:

            if "@nome" in paragrafo.text:

                paragrafo.text = nome_separado
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@DataFim" in paragrafo.text:

                paragrafo.text = frase_montada
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@instrutor" in paragrafo.text:

                paragrafo.text = f"{nomeInstrutor} - Instrutor"
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        #Montando o caminho + nome do certificado
        caminhoCertificadoGerado = nome_separado + ".docx"

        #Salvando o certificado com o nome do aluno
        arquivoWord.save(caminhoCertificadoGerado)

    messagebox.showinfo("Mensagem", "Certificados gerados com sucesso!")
#---------------------------------------------------------------------

botaoPesquisar = Button(text="PESQUISAR", font="Arial 10 bold", command=filtrar_dados)
botaoPesquisar.grid(row=5, column=0, columnspan=3, sticky="NSEW", padx=10)

botaoGerarCertificado = Button(text="GERAR CERTIFICADO", font="Arial 10 bold", command=gerar_certificado)
botaoGerarCertificado.grid(row=5, column=3, columnspan=3, sticky="NSEW", padx=10)

botaoGerarCertificadoMassa = Button(text="GERAR EM MASSA", font="Arial 10 bold", command=gerar_certificado_massa)
botaoGerarCertificadoMassa.grid(row=5, column=6, columnspan=4, sticky="NSEW", padx=10)


root.mainloop()